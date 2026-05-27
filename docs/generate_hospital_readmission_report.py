from __future__ import annotations

import json
import math
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs" / "22636801_TranThaiHa.docx"
OUT = ROOT / "docs" / "Hospital_Readmission_Prediction_Report.docx"
DIAGRAM_DIR = ROOT / "docs" / "generated_diagrams"
SOURCE_IMAGE_DIR = ROOT / "docs" / "images"
METRICS_PATH = ROOT / "reports" / "metrics.json"
LOGO_PATH = DIAGRAM_DIR / "iuh_logo.png"


SOURCE_DIAGRAMS = {
    "function": "ChatGPT Image 02_11_32 24 thg 5, 2026 (1).png",
    "use_case": "ChatGPT Image 02_11_32 24 thg 5, 2026 (2).png",
    "activity": "ChatGPT Image 02_11_32 24 thg 5, 2026 (3).png",
    "sequence": "ChatGPT Image 02_11_32 24 thg 5, 2026 (4).png",
    "class": "ChatGPT Image 02_11_33 24 thg 5, 2026 (5).png",
    "dfd": "ChatGPT Image 02_11_33 24 thg 5, 2026 (6).png",
    "erd": "ChatGPT Image 02_11_35 24 thg 5, 2026 (7).png",
}


COLORS = {
    "blue": "#1f5fbf",
    "dark_blue": "#0f2b5f",
    "green": "#237a3b",
    "orange": "#d66a1f",
    "red": "#b4232a",
    "purple": "#6a3aa8",
    "teal": "#13827f",
    "yellow": "#c08a16",
    "gray": "#8a8f98",
    "light": "#f3f7fb",
    "line": "#1f2937",
}


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


FONT_17 = font(17)
FONT_18 = font(18)
FONT_20 = font(20)
FONT_22 = font(22)
FONT_24_B = font(24, True)
FONT_28_B = font(28, True)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=fnt)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if text_size(draw, candidate, fnt)[0] <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fnt: ImageFont.ImageFont,
    fill: str = "#ffffff",
    line_spacing: int = 8,
) -> None:
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 28
    lines = wrap_text(draw, text, fnt, max_width)
    heights = [text_size(draw, line, fnt)[1] for line in lines]
    total = sum(heights) + line_spacing * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, height in zip(lines, heights):
        width = text_size(draw, line, fnt)[0]
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=fnt, fill=fill)
        y += height + line_spacing


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str = "#202020",
    text_fill: str = "#ffffff",
    radius: int = 18,
    fnt: ImageFont.ImageFont = FONT_20,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=2)
    draw_centered_text(draw, box, text, fnt, text_fill)


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: str = "#202020",
    width: int = 4,
) -> None:
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        points = [(ex, ey), (ex - direction * 16, ey - 9), (ex - direction * 16, ey + 9)]
    else:
        direction = 1 if ey >= sy else -1
        points = [(ex, ey), (ex - 9, ey - direction * 16), (ex + 9, ey - direction * 16)]
    draw.polygon(points, fill=fill)


def dashed_line(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: str = "#202020",
    width: int = 3,
    dash: int = 16,
    gap: int = 10,
) -> None:
    x1, y1 = start
    x2, y2 = end
    dx = x2 - x1
    dy = y2 - y1
    distance = max((dx * dx + dy * dy) ** 0.5, 1)
    step = dash + gap
    t = 0.0
    while t < distance:
        t2 = min(t + dash, distance)
        draw.line(
            (
                x1 + dx * t / distance,
                y1 + dy * t / distance,
                x1 + dx * t2 / distance,
                y1 + dy * t2 / distance,
            ),
            fill=fill,
            width=width,
        )
        t += step


def new_canvas(title: str, size: tuple[int, int] = (1600, 950)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, "#ffffff")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, size[0], 72), fill=COLORS["dark_blue"])
    draw.text((40, 19), title, font=FONT_28_B, fill="#ffffff")
    return image, draw


def save(image: Image.Image, name: str) -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / name
    image.save(path, quality=95)
    return path


def prepare_source_diagram(key: str, output_name: str) -> Path:
    source_path = SOURCE_IMAGE_DIR / SOURCE_DIAGRAMS[key]
    image = Image.open(source_path).convert("RGB")
    return save(image, output_name)


def extract_logo() -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(TEMPLATE) as docx:
        LOGO_PATH.write_bytes(docx.read("word/media/image1.png"))
    return LOGO_PATH


def draw_function_diagram() -> Path:
    image = Image.new("RGB", (1620, 1700), "#ffffff")
    draw = ImageDraw.Draw(image)

    root = (380, 55, 1240, 155)
    draw_box(
        draw,
        root,
        "Hệ thống dự đoán tái nhập viện bệnh nhân",
        "#eef6ff",
        outline=COLORS["line"],
        text_fill=COLORS["line"],
        radius=16,
        fnt=FONT_24_B,
    )

    groups = [
        (
            "1. Quản lý\nngười dùng và\nxác thực",
            COLORS["line"],
            "#eef6ff",
            [
                "Đăng nhập",
                "Cấp JWT token",
                "Phân quyền Doctor/Admin",
                "Xem thông tin người dùng",
            ],
        ),
        (
            "2. Quản lý\nbệnh nhân",
            COLORS["line"],
            "#f0fbf0",
            [
                "Thêm hồ sơ bệnh nhân",
                "Xem danh sách bệnh nhân",
                "Xem chi tiết bệnh nhân",
                "Cập nhật hồ sơ bệnh nhân",
                "Xóa hồ sơ bệnh nhân",
                "Xem bệnh nhân nguy cơ cao",
            ],
        ),
        (
            "3. Dự đoán tái\nnhập viện",
            COLORS["line"],
            "#fff7dc",
            [
                "Nhập dữ liệu lâm sàng",
                "Dự đoán từ hồ sơ bệnh nhân",
                "Tiền xử lý dữ liệu",
                "Tải model Production",
                "Trả kết quả dự đoán",
                "Ghi log dự đoán",
            ],
        ),
        (
            "4. Quản trị\nMLOps",
            COLORS["line"],
            "#f7f0ff",
            [
                "Xem trạng thái pipeline",
                "Kích hoạt Airflow DAG",
                "Chạy ETL",
                "Huấn luyện model",
                "Retraining model",
                "Reload model Production",
            ],
        ),
        (
            "5. Quản lý dữ liệu\nvà huấn luyện\nmô hình",
            COLORS["line"],
            "#fff0dc",
            [
                "Đọc dữ liệu nguồn",
                "Spark Batch ETL",
                "Tạo dữ liệu Gold",
                "Tạo Offline Feature",
                "Train model",
                "Đánh giá model",
                "Ghi experiment vào MLflow",
                "Lưu artifact vào MinIO",
                "Promote model lên Production",
            ],
        ),
        (
            "6. Giám sát\nhệ thống",
            COLORS["line"],
            "#ffe7e7",
            [
                "Thu thập metrics từ FastAPI",
                "Prometheus scrape metrics",
                "Grafana dashboard",
                "Cảnh báo hệ thống",
                "Theo dõi request, latency, error và risk level",
            ],
        ),
    ]

    column_centers = [135, 405, 675, 945, 1215, 1485]
    header_y = 280
    header_size = (220, 135)
    item_size = (190, 90)
    item_gap = 40
    branch_y = 220
    root_center_x = (root[0] + root[2]) // 2

    draw.line((root_center_x, root[3], root_center_x, branch_y), fill=COLORS["line"], width=3)
    draw.line((column_centers[0], branch_y, column_centers[-1], branch_y), fill=COLORS["line"], width=3)

    for center_x, (title, color, fill, items) in zip(column_centers, groups):
        header = (
            center_x - header_size[0] // 2,
            header_y,
            center_x + header_size[0] // 2,
            header_y + header_size[1],
        )
        arrow(draw, (center_x, branch_y), (center_x, header[1]), fill=color, width=3)
        draw_box(
            draw,
            header,
            title,
            fill,
            outline=color,
            text_fill=COLORS["line"],
            radius=14,
            fnt=FONT_22,
        )

        line_x = header[0] + 10
        line_top = header[3]
        line_bottom = header[3] + len(items) * item_size[1] + (len(items) - 1) * item_gap + 35
        draw.line((line_x, line_top, line_x, line_bottom), fill=color, width=3)

        y = header[3] + 55
        for item in items:
            item_box = (
                center_x - item_size[0] // 2,
                y,
                center_x + item_size[0] // 2,
                y + item_size[1],
            )
            arrow(draw, (line_x, y + item_size[1] // 2), (item_box[0], y + item_size[1] // 2), fill=color, width=2)
            draw_box(
                draw,
                item_box,
                item,
                "#ffffff",
                outline=color,
                text_fill=COLORS["line"],
                radius=12,
                fnt=FONT_17,
            )
            y += item_size[1] + item_gap

    return save(image, "01_function_diagram.png")


def draw_use_case() -> Path:
    image = Image.new("RGB", (1700, 1250), "#ffffff")
    draw = ImageDraw.Draw(image)

    def actor(cx: int, y: int, label: str) -> None:
        draw.ellipse((cx - 28, y, cx + 28, y + 56), fill="#ffffff", outline="#111111", width=3)
        draw.line((cx, y + 56, cx, y + 165), fill="#111111", width=4)
        draw.line((cx - 58, y + 105, cx + 58, y + 105), fill="#111111", width=4)
        draw.line((cx, y + 165, cx - 45, y + 265), fill="#111111", width=4)
        draw.line((cx, y + 165, cx + 45, y + 265), fill="#111111", width=4)
        tw, _ = text_size(draw, label, FONT_24_B)
        draw.text((cx - tw / 2, y + 288), label, font=FONT_24_B, fill="#111111")

    def use_case(box: tuple[int, int, int, int], label: str) -> None:
        draw.ellipse(box, fill="#ffffff", outline="#111111", width=3)
        draw_centered_text(draw, box, label, FONT_18, "#111111")

    def association(start: tuple[int, int], end: tuple[int, int]) -> None:
        draw.line((start, end), fill="#111111", width=3)

    def dashed_connector(
        start: tuple[int, int],
        end: tuple[int, int],
        label: str = "",
        label_offset: tuple[int, int] = (0, 0),
        arrow_head: bool = True,
    ) -> None:
        x1, y1 = start
        x2, y2 = end
        dashed_line(draw, start, end, "#111111", width=2, dash=14, gap=8)
        if arrow_head:
            angle = math.atan2(y2 - y1, x2 - x1)
            head_len = 18
            head_angle = math.radians(28)
            for side in (-1, 1):
                hx = x2 - head_len * math.cos(angle + side * head_angle)
                hy = y2 - head_len * math.sin(angle + side * head_angle)
                draw.line((x2, y2, hx, hy), fill="#111111", width=2)
        if label:
            mx = (x1 + x2) / 2 + label_offset[0]
            my = (y1 + y2) / 2 + label_offset[1]
            tw, th = text_size(draw, label, FONT_18)
            draw.rectangle((mx - tw / 2 - 8, my - th / 2 - 5, mx + tw / 2 + 8, my + th / 2 + 5), fill="#ffffff")
            draw.text((mx - tw / 2, my - th / 2), label, font=FONT_18, fill="#111111")

    boundary = (220, 25, 1480, 1205)
    draw.rectangle(boundary, outline="#111111", width=4)

    title = "Hospital Readmission Prediction MLOps"
    title_width, _ = text_size(draw, title, FONT_28_B)
    draw.text((850 - title_width / 2, 48), title, font=FONT_28_B, fill="#111111")

    actor(88, 420, "Doctor")
    actor(1612, 420, "Admin")

    doctor_cases = {
        "login": ((285, 110, 530, 205), "Đăng nhập"),
        "profile": ((285, 245, 540, 335), "Xem thông tin cá nhân"),
        "patients": ((285, 390, 565, 485), "Quản lý hồ sơ bệnh nhân"),
        "predict_direct": ((275, 535, 575, 635), "Dự đoán từ form lâm sàng"),
        "predict_saved": ((275, 695, 575, 795), "Dự đoán từ hồ sơ đã lưu"),
        "history": ((285, 865, 565, 960), "Xem lịch sử dự đoán"),
        "highrisk": ((285, 1010, 565, 1105), "Xem bệnh nhân nguy cơ cao"),
    }
    admin_cases = {
        "login": ((1145, 110, 1390, 205), "Đăng nhập"),
        "dashboard": ((1130, 245, 1410, 335), "Xem dashboard thống kê"),
        "all_patients": ((1130, 390, 1410, 485), "Xem toàn bộ bệnh nhân"),
        "logs": ((1130, 535, 1410, 635), "Xem toàn bộ\nprediction logs"),
        "users": ((1130, 695, 1410, 795), "Quản lý người dùng"),
        "status": ((1130, 850, 1410, 950), "Bật / tắt tài khoản"),
        "pipelines": ((1120, 990, 1420, 1090), "Theo dõi / kích hoạt pipeline"),
        "reload": ((1120, 1120, 1420, 1200), "Reload model Production"),
    }
    internal_cases = {
        "auth": ((725, 160, 975, 250), "Xác thực JWT"),
        "preprocess": ((725, 385, 975, 480), "Tiền xử lý dữ liệu"),
        "model": ((725, 535, 975, 635), "Tải model Production"),
        "log": ((725, 695, 975, 795), "Ghi prediction log"),
        "airflow": ((725, 860, 975, 960), "Airflow API"),
        "registry": ((725, 1030, 975, 1130), "MLflow Registry"),
    }

    for box, label in doctor_cases.values():
        use_case(box, label)
    for box, label in admin_cases.values():
        use_case(box, label)
    for box, label in internal_cases.values():
        use_case(box, label)

    for box, _ in doctor_cases.values():
        association((148, 555), (box[0], (box[1] + box[3]) // 2))
    for box, _ in admin_cases.values():
        association((1552, 555), (box[2], (box[1] + box[3]) // 2))

    def include(start: tuple[int, int], end: tuple[int, int], label_offset: tuple[int, int]) -> None:
        dashed_connector(start, end, "<<include>>", label_offset)

    def extend(start: tuple[int, int], end: tuple[int, int], label_offset: tuple[int, int]) -> None:
        dashed_connector(start, end, "<<extend>>", label_offset)

    # Include: use case always reuses a common service.
    include((530, 158), (725, 205), (10, -42))
    include((1145, 158), (975, 205), (-12, -42))
    include((575, 585), (725, 430), (20, -25))
    include((575, 585), (725, 585), (-70, -35))
    include((575, 745), (725, 745), (-70, 34))
    include((1120, 1040), (975, 912), (78, -8))
    include((1120, 1160), (975, 1080), (0, -42))

    # Extend: optional monitoring or management action extends a base use case.
    extend((1130, 585), (975, 745), (76, -2))
    extend((565, 912), (725, 745), (55, 15))
    extend((1270, 850), (1270, 795), (112, -4))

    return save(image, "02_use_case.png")

def draw_activity() -> Path:
    image, draw = new_canvas("Biểu đồ hoạt động dự đoán và ghi nhận phản hồi")
    steps = [
        ("Bác sĩ đăng nhập", COLORS["blue"]),
        ("Tạo/cập nhật hồ sơ bệnh nhân", COLORS["blue"]),
        ("Gửi yêu cầu dự đoán", COLORS["green"]),
        ("FastAPI chuẩn hóa đặc trưng", COLORS["green"]),
        ("Model MLflow trả xác suất", COLORS["orange"]),
        ("Phân mức low / medium / high", COLORS["orange"]),
        ("Lưu log vào Redis và PostgreSQL", COLORS["purple"]),
        ("Cập nhật dashboard / danh sách high-risk", COLORS["purple"]),
    ]
    y = 120
    previous = None
    for label, color in steps:
        box = (530, y, 1070, y + 65)
        draw_box(draw, box, label, color, fnt=FONT_20)
        if previous:
            arrow(draw, ((previous[0] + previous[2]) // 2, previous[3]), ((box[0] + box[2]) // 2, box[1]))
        previous = box
        y += 92
    decision = (1120, 477, 1460, 577)
    draw.polygon([(1290, 440), (1465, 527), (1290, 615), (1115, 527)], fill="#f5f5f5", outline=COLORS["orange"])
    draw_centered_text(draw, (1135, 460, 1445, 595), "Đủ nhãn mới\nđể tái huấn luyện?", FONT_18, "#202020")
    arrow(draw, (1070, 520), (1115, 527), fill=COLORS["orange"])
    draw_box(draw, (1120, 660, 1460, 725), "Airflow retraining DAG", COLORS["orange"], fnt=FONT_18)
    arrow(draw, (1290, 615), (1290, 660), fill=COLORS["orange"])
    draw.text((1320, 625), "Có", font=FONT_18, fill=COLORS["orange"])
    draw_box(draw, (1120, 775, 1460, 840), "Đăng ký model mới và reload", COLORS["red"], fnt=FONT_18)
    arrow(draw, (1290, 725), (1290, 775), fill=COLORS["red"])
    draw.text((1468, 526), "Chưa đủ: tiếp tục thu thập", font=FONT_18, fill=COLORS["gray"])
    return save(image, "03_activity.png")


def draw_sequence() -> Path:
    image, draw = new_canvas("Biểu đồ trình tự xử lý một yêu cầu dự đoán")
    actors = ["React UI", "FastAPI", "ModelLoader", "MLflow Registry", "Redis", "PostgreSQL", "Prometheus"]
    xs = [130, 360, 590, 820, 1040, 1260, 1470]
    for x, actor in zip(xs, actors):
        draw_box(draw, (x - 80, 120, x + 80, 175), actor, COLORS["dark_blue"], fnt=FONT_18)
        draw.line((x, 175, x, 835), fill="#bdbdbd", width=3)
    messages = [
        (0, 1, 230, "POST /patients/{id}/predict"),
        (1, 2, 310, "predict_probability(df)"),
        (2, 3, 390, "load models:/Production nếu cần"),
        (3, 2, 470, "sklearn pipeline"),
        (2, 1, 550, "probability"),
        (1, 4, 630, "cache prediction log"),
        (1, 5, 710, "insert prediction_logs"),
        (1, 6, 790, "observe metrics"),
    ]
    for a, b, y, msg in messages:
        arrow(draw, (xs[a] + 82, y), (xs[b] - 82 if b > a else xs[b] + 82, y), width=3)
        draw.text((min(xs[a], xs[b]) + 25, y - 28), msg, font=FONT_18, fill=COLORS["line"])
    arrow(draw, (360, 835), (130, 835), width=3)
    draw.text((160, 806), "PredictionResponse: prediction, probability, risk_level", font=FONT_18, fill=COLORS["line"])
    return save(image, "04_sequence.png")


def draw_class_diagram() -> Path:
    image, draw = new_canvas("Biểu đồ lớp - Hospital Readmission Prediction MLOps", (1800, 1350))

    def class_box(
        box: tuple[int, int, int, int],
        name: str,
        attrs: list[str],
        methods: list[str],
        color: str,
    ) -> None:
        x1, y1, x2, y2 = box
        draw.rounded_rectangle(box, radius=0, fill="#ffffff", outline=color, width=4)
        draw.rectangle((x1, y1, x2, y1 + 52), fill="#eef4ff", outline=color, width=3)
        draw.text((x1 + 18, y1 + 13), name, font=FONT_22, fill=color)
        divider = y1 + 52 + max(130, len(attrs) * 30 + 18)
        draw.line((x1, divider, x2, divider), fill=color, width=3)
        y = y1 + 70
        for attr in attrs:
            draw.text((x1 + 18, y), attr, font=FONT_18, fill=COLORS["line"])
            y += 30
        y = divider + 18
        for method in methods:
            draw.text((x1 + 18, y), method, font=FONT_18, fill=COLORS["line"])
            y += 30

    class_box(
        (60, 130, 380, 570),
        "User",
        ["- id: int PK", "- username: str", "- password_hash: str", "- full_name: str", "- role: doctor/admin", "- is_active: bool", "- created_at: datetime"],
        ["+ authenticate()", "+ get_current_user()", "+ require_admin()", "+ set_active()"],
        COLORS["blue"],
    )
    class_box(
        (500, 220, 850, 680),
        "Patient",
        ["- id: int PK", "- doctor_id: int FK", "- demographics", "- admission fields", "- diagnosis fields", "- medication fields", "- actual_readmitted: int?", "- created_at/updated_at"],
        ["+ create_patient()", "+ update_patient()", "+ list_by_doctor()", "+ get_high_risk()"],
        COLORS["green"],
    )
    class_box(
        (965, 220, 1340, 690),
        "PredictionLog",
        ["- id: int PK", "- patient_id: int? FK", "- doctor_id: int? FK", "- request_json: JSONB", "- prediction: int", "- probability: float", "- risk_level: str", "- model_version/run_id", "- created_at: datetime"],
        ["+ save_prediction_log()", "+ list_logs()", "+ list_by_patient()"],
        COLORS["orange"],
    )
    class_box(
        (1425, 150, 1740, 445),
        "RetrainingState",
        ["- id: int PK", "- last_trained_prediction_count", "- last_trained_patient_count", "- updated_at: datetime"],
        ["+ read_state()", "+ update_state()"],
        COLORS["purple"],
    )
    class_box(
        (1425, 560, 1740, 1010),
        "RetrainingRun",
        ["- id: int PK", "- trigger_type: str", "- new_records: int", "- status: str", "- metric_name/value", "- started_at/ended_at", "- created_at: datetime"],
        ["+ create_run()", "+ mark_success()", "+ mark_skipped()"],
        COLORS["red"],
    )
    class_box(
        (90, 890, 500, 1245),
        "AuthService",
        ["+ hash/verify password", "+ create JWT", "+ resolve role", "+ reject disabled account"],
        ["+ login()", "+ require_doctor_or_admin()"],
        COLORS["yellow"],
    )
    class_box(
        (650, 890, 1090, 1275),
        "PredictionService",
        ["+ normalize payload", "+ build dataframe", "+ call ModelLoader", "+ classify risk"],
        ["+ predict_direct()", "+ predict_patient()", "+ persist metrics/logs()"],
        COLORS["teal"],
    )
    class_box(
        (1260, 1060, 1665, 1330),
        "AirflowService",
        ["+ allowed DAG configs", "+ Airflow REST client", "+ run status mapping"],
        ["+ get_pipeline_status()", "+ trigger_dag()"],
        COLORS["dark_blue"],
    )

    arrow(draw, (380, 315), (500, 315), fill=COLORS["line"], width=3)
    draw.text((410, 285), "1 - N", font=FONT_18, fill=COLORS["line"])
    arrow(draw, (850, 430), (965, 430), fill=COLORS["line"], width=3)
    draw.text((882, 400), "1 - N", font=FONT_18, fill=COLORS["line"])
    arrow(draw, (380, 210), (965, 270), fill=COLORS["line"], width=3)
    draw.text((650, 195), "doctor creates logs", font=FONT_18, fill=COLORS["line"])
    dashed_line(draw, (295, 890), (250, 570), fill=COLORS["yellow"], width=3)
    dashed_line(draw, (870, 890), (675, 680), fill=COLORS["teal"], width=3)
    dashed_line(draw, (870, 890), (1150, 690), fill=COLORS["teal"], width=3)
    dashed_line(draw, (1260, 1100), (1425, 930), fill=COLORS["dark_blue"], width=3)
    dashed_line(draw, (1665, 1100), (1585, 445), fill=COLORS["dark_blue"], width=3)
    draw.text((1305, 1015), "orchestrates", font=FONT_18, fill=COLORS["line"])

    return save(image, "05_class_diagram.png")


def draw_erd() -> Path:
    image, draw = new_canvas("Biểu đồ quan hệ dữ liệu PostgreSQL", (1800, 1250))
    tables = {
        "users": ((70, 135, 390, 420), ["id PK", "username UNIQUE", "password_hash", "full_name", "role", "is_active", "created_at"]),
        "patients": ((620, 120, 1050, 610), ["id PK", "doctor_id FK", "race, gender, age", "admission_type_id", "discharge_disposition_id", "admission_source_id", "time_in_hospital", "lab/procedure/med counts", "diagnoses", "medication flags", "actual_readmitted", "created_at, updated_at"]),
        "prediction_logs": ((620, 745, 1050, 1165), ["id PK", "patient_id FK nullable", "doctor_id FK nullable", "request_json JSONB", "prediction", "readmission_probability", "risk_level", "model_name", "model_version", "model_run_id", "created_at"]),
        "retraining_state": ((1300, 160, 1690, 390), ["id PK", "last_trained_prediction_count", "last_trained_patient_count", "updated_at"]),
        "retraining_runs": ((1300, 650, 1690, 1010), ["id PK", "trigger_type", "new_records", "status", "metric_name", "metric_value", "started_at", "ended_at", "created_at"]),
    }
    table_colors = {
        "users": COLORS["blue"],
        "patients": COLORS["green"],
        "prediction_logs": COLORS["orange"],
        "retraining_state": COLORS["purple"],
        "retraining_runs": COLORS["yellow"],
    }
    for name, (box, fields) in tables.items():
        color = table_colors[name]
        draw.rounded_rectangle(box, radius=8, fill="#ffffff", outline=color, width=3)
        draw.rectangle((box[0], box[1], box[2], box[1] + 48), fill=color)
        draw.text((box[0] + 14, box[1] + 10), name, font=FONT_20, fill="#ffffff")
        y = box[1] + 62
        for field in fields:
            draw.text((box[0] + 18, y), field, font=FONT_18, fill=COLORS["line"])
            y += 28

    draw.line((390, 260, 620, 260), fill=COLORS["line"], width=4)
    draw.text((435, 225), "1", font=FONT_24_B, fill=COLORS["line"])
    draw.text((575, 225), "N", font=FONT_24_B, fill=COLORS["line"])
    draw.text((425, 285), "doctor_id", font=FONT_18, fill=COLORS["line"])

    draw.line((835, 610, 835, 745), fill=COLORS["line"], width=4)
    draw.text((805, 625), "1", font=FONT_24_B, fill=COLORS["line"])
    draw.text((805, 705), "N", font=FONT_24_B, fill=COLORS["line"])

    draw.line((230, 390, 230, 915), fill=COLORS["line"], width=4)
    draw.line((230, 915, 620, 915), fill=COLORS["line"], width=4)
    draw.text((245, 420), "1", font=FONT_24_B, fill=COLORS["line"])
    draw.text((575, 880), "N", font=FONT_24_B, fill=COLORS["line"])
    draw.text((265, 890), "doctor_id", font=FONT_18, fill=COLORS["line"])

    dashed_line(draw, (1495, 390), (1495, 650), fill=COLORS["purple"], width=4)
    draw.text((1515, 500), "logical run history", font=FONT_18, fill=COLORS["purple"])

    legend = (70, 920, 390, 1040)
    draw.rectangle(legend, fill="#ffffff", outline=COLORS["line"], width=2)
    draw.text((92, 942), "PK: Primary Key", font=FONT_18, fill=COLORS["line"])
    draw.text((92, 978), "FK: Foreign Key", font=FONT_18, fill=COLORS["line"])
    draw.text((92, 1014), "Nét đứt: quan hệ logic", font=FONT_18, fill=COLORS["line"])
    return save(image, "06_erd.png")


def draw_dfd() -> Path:
    image, draw = new_canvas("Biểu đồ luồng dữ liệu batch, streaming và phục vụ model")
    nodes = [
        ((60, 150, 330, 235), "CSV diabetic_data.csv", COLORS["gray"]),
        ((455, 150, 735, 235), "Spark batch ETL", COLORS["orange"]),
        ((860, 150, 1135, 235), "Gold parquet", COLORS["green"]),
        ((1245, 150, 1535, 235), "Offline feature parquet", COLORS["green"]),
        ((1245, 360, 1535, 445), "Training pipeline", COLORS["purple"]),
        ((860, 360, 1135, 445), "MLflow Registry\nProduction model", COLORS["purple"]),
        ((455, 360, 735, 445), "FastAPI inference", COLORS["red"]),
        ((60, 360, 330, 445), "React UI\nDoctor/Admin", COLORS["blue"]),
        ((455, 600, 735, 685), "PostgreSQL + Redis\nprediction logs", COLORS["orange"]),
        ((860, 600, 1135, 685), "Airflow DB retraining", COLORS["dark_blue"]),
        ((455, 815, 735, 890), "Kafka patient-events", COLORS["gray"]),
        ((860, 815, 1135, 890), "Spark Streaming\noptional profile", COLORS["gray"]),
        ((1245, 815, 1535, 890), "Bronze parquet\noptional audit/demo", COLORS["gray"]),
    ]
    for box, label, color in nodes:
        draw_box(draw, box, label, color, fnt=FONT_18)

    arrow(draw, (330, 192), (455, 192))
    arrow(draw, (735, 192), (860, 192))
    arrow(draw, (1135, 192), (1245, 192))
    arrow(draw, (1390, 235), (1390, 360), fill=COLORS["purple"])
    arrow(draw, (1245, 402), (1135, 402), fill=COLORS["purple"])
    arrow(draw, (860, 402), (735, 402), fill=COLORS["purple"])
    arrow(draw, (330, 402), (455, 402), fill=COLORS["blue"])
    arrow(draw, (595, 445), (595, 600), fill=COLORS["orange"])
    arrow(draw, (735, 642), (860, 642), fill=COLORS["orange"])
    draw.line((1135, 642, 1390, 642), fill=COLORS["dark_blue"], width=4)
    draw.line((1390, 642, 1390, 445), fill=COLORS["dark_blue"], width=4)
    arrow(draw, (1390, 642), (1390, 445), fill=COLORS["dark_blue"])
    arrow(draw, (735, 852), (860, 852), fill=COLORS["gray"])
    arrow(draw, (1135, 852), (1245, 852), fill=COLORS["gray"])
    draw.text((300, 285), "Airflow điều phối ingestion, ETL, training và DB-triggered retraining", font=FONT_24_B, fill=COLORS["dark_blue"])
    draw.text((440, 775), "Luồng streaming là tùy chọn, không dùng làm trigger retraining tự động", font=FONT_20, fill=COLORS["line"])
    return save(image, "07_dfd.png")


def draw_algorithm() -> Path:
    image, draw = new_canvas("Quy trình tiền xử lý, huấn luyện và lựa chọn mô hình")
    boxes = [
        ((75, 195, 340, 285), "Đọc CSV\n101.766 dòng, 50 cột", COLORS["gray"]),
        ((430, 195, 695, 285), "Làm sạch\n? -> missing, drop cột yếu", COLORS["blue"]),
        ((785, 195, 1050, 285), "Tạo target\nreadmitted = <30", COLORS["green"]),
        ((1140, 195, 1505, 285), "Chia train/test\nstratify, test_size=0.2", COLORS["orange"]),
        ((145, 475, 450, 565), "ColumnTransformer\nimpute + scale + one-hot", COLORS["purple"]),
        ((550, 475, 855, 565), "Ứng viên 1\nRandom Forest", COLORS["green"]),
        ((955, 475, 1260, 565), "Ứng viên 2\nLogistic Regression", COLORS["blue"]),
        ((590, 710, 1010, 805), "Chọn champion theo ROC-AUC\nlog MLflow, lưu model.pkl", COLORS["red"]),
    ]
    for box, label, color in boxes:
        draw_box(draw, box, label, color, fnt=FONT_18)
    for i in range(3):
        arrow(draw, (boxes[i][0][2], 240), (boxes[i + 1][0][0], 240))
    arrow(draw, (1320, 285), (300, 475))
    arrow(draw, (450, 520), (550, 520))
    arrow(draw, (855, 520), (955, 520))
    arrow(draw, (700, 565), (700, 710), fill=COLORS["red"])
    arrow(draw, (1110, 565), (890, 710), fill=COLORS["red"])
    draw.text((390, 350), "Metrics: accuracy, precision, recall, f1_score, roc_auc", font=FONT_22, fill=COLORS["line"])
    return save(image, "08_algorithm.png")


def draw_retraining() -> Path:
    image, draw = new_canvas("Luồng tái huấn luyện tự động")
    steps = [
        ((70, 155, 355, 235), "Bác sĩ/Admin\ncập nhật actual_readmitted", COLORS["blue"]),
        ((445, 155, 730, 235), "PostgreSQL\npatients + prediction_logs", COLORS["green"]),
        ((820, 155, 1105, 235), "Airflow\nDB-triggered DAG", COLORS["dark_blue"]),
        ((1210, 155, 1510, 235), "Đủ nhãn mới?\n>= RETRAIN_MIN_NEW_LABELS", COLORS["gray"]),
        ((1210, 340, 1510, 420), "Skip\nchờ thêm nhãn", COLORS["gray"]),
        ((820, 340, 1105, 420), "prepare_training_from_db\n-> db_training_data.csv", COLORS["green"]),
        ((820, 520, 1105, 600), "training.train\ncandidate models", COLORS["orange"]),
        ((430, 520, 715, 600), "MLflow runs\nmetrics + artifacts", COLORS["purple"]),
        ((70, 520, 355, 600), "register_model\nregister v1, v2...", COLORS["purple"]),
        ((430, 735, 715, 815), "FastAPI /reload-model\nModelLoader refresh", COLORS["red"]),
        ((820, 735, 1105, 815), "Update retraining_state\nvà retraining_runs", COLORS["dark_blue"]),
    ]
    for box, label, color in steps:
        draw_box(draw, box, label, color, fnt=FONT_18)

    arrow(draw, (355, 195), (445, 195))
    arrow(draw, (730, 195), (820, 195))
    arrow(draw, (1105, 195), (1210, 195))
    arrow(draw, (1360, 235), (1360, 340), fill=COLORS["gray"])
    arrow(draw, (1210, 235), (1105, 380), fill=COLORS["green"])
    arrow(draw, (962, 420), (962, 520), fill=COLORS["orange"])
    arrow(draw, (820, 560), (715, 560), fill=COLORS["purple"])
    arrow(draw, (430, 560), (355, 560), fill=COLORS["purple"])
    draw.line((212, 600, 212, 775), fill=COLORS["red"], width=4)
    arrow(draw, (212, 775), (430, 775), fill=COLORS["red"])
    arrow(draw, (715, 775), (820, 775), fill=COLORS["dark_blue"])

    draw.text((1385, 290), "Không", font=FONT_20, fill=COLORS["line"])
    draw.text((1130, 310), "Có", font=FONT_20, fill=COLORS["line"])
    return save(image, "09_retraining.png")


def draw_mlops_diagram() -> Path:
    image, draw = new_canvas("Sơ đồ MLOps tổng thể")

    lifecycle = [
        ((70, 155, 330, 240), "Data Source\nCSV, Kafka, PostgreSQL", COLORS["gray"]),
        ((430, 155, 690, 240), "Data Pipeline\nSpark ETL, Feature Engineering", COLORS["blue"]),
        ((790, 155, 1050, 240), "Experiment Tracking\nMLflow Runs, Metrics", COLORS["purple"]),
        ((1150, 155, 1500, 240), "Model Registry\nProduction Version", COLORS["dark_blue"]),
        ((1150, 410, 1500, 495), "Serving\nFastAPI + ModelLoader", COLORS["red"]),
        ((790, 410, 1050, 495), "Application\nReact UI + JWT API", COLORS["green"]),
        ((430, 410, 690, 495), "Prediction Logs\nPostgreSQL + Redis", COLORS["orange"]),
        ((70, 410, 330, 495), "Monitoring\nPrometheus + Grafana", COLORS["teal"]),
        ((610, 675, 990, 770), "Orchestration\nAirflow DAGs: ingestion, ETL,\ntraining, DB retraining, reload", COLORS["dark_blue"]),
    ]

    for box, label, color in lifecycle:
        draw_box(draw, box, label, color, fnt=FONT_18)

    for start, end in [
        ((330, 197), (430, 197)),
        ((690, 197), (790, 197)),
        ((1050, 197), (1150, 197)),
        ((1325, 240), (1325, 410)),
        ((1150, 452), (1050, 452)),
        ((790, 452), (690, 452)),
        ((430, 452), (330, 452)),
        ((200, 410), (200, 240)),
    ]:
        arrow(draw, start, end)

    draw.line((800, 675, 800, 585), fill=COLORS["line"], width=3)
    draw.line((800, 585, 560, 585), fill=COLORS["line"], width=3)
    arrow(draw, (560, 585), (560, 495), fill=COLORS["line"], width=3)
    draw.line((990, 722, 1325, 722), fill=COLORS["line"], width=3)
    draw.line((1325, 722, 1325, 495), fill=COLORS["line"], width=3)
    arrow(draw, (1325, 722), (1325, 495), fill=COLORS["line"], width=3)
    draw.line((610, 722, 200, 722), fill=COLORS["line"], width=3)
    draw.line((200, 722, 200, 495), fill=COLORS["line"], width=3)
    arrow(draw, (200, 722), (200, 495), fill=COLORS["line"], width=3)

    draw.rectangle((360, 295, 1240, 340), fill="#ffffff")
    draw.text((410, 305), "Champion model được chọn theo ROC-AUC và promote khi tốt hơn Production", font=FONT_20, fill=COLORS["line"])
    draw.rectangle((370, 550, 1225, 595), fill="#ffffff")
    draw.text((420, 560), "Dữ liệu vận hành và nhãn mới quay lại pipeline để tái huấn luyện", font=FONT_20, fill=COLORS["line"])
    return save(image, "10_mlops_diagram.png")


def draw_deployment() -> Path:
    image, draw = new_canvas("Sơ đồ triển khai Docker Compose")
    groups = [
        ("Data & Streaming", 60, 130, 475, 410, [("Kafka", "Zookeeper", "Kafka UI"), ("Spark Master", "Spark Worker", "Streaming profile")], COLORS["blue"]),
        ("ML & Storage", 555, 130, 955, 410, [("MLflow", "MinIO", "Trainer profile"), ("Models", "Reports", "DVC")], COLORS["purple"]),
        ("Application", 1035, 130, 1535, 410, [("FastAPI", "PostgreSQL", "Redis"), ("React Frontend", "Nginx", "Auth/JWT")], COLORS["green"]),
        ("Orchestration", 170, 555, 690, 820, [("Airflow Webserver", "Airflow Scheduler", "DAGs"), ("Ingestion", "ETL", "DB Retraining")], COLORS["orange"]),
        ("Monitoring", 880, 555, 1430, 820, [("Prometheus", "Grafana", "Alert rules"), ("API metrics", "Latency", "Errors")], COLORS["red"]),
    ]
    for title, x1, y1, x2, y2, rows, color in groups:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill="#f7f7f7", outline=color, width=4)
        draw.text((x1 + 24, y1 + 20), title, font=FONT_24_B, fill=color)
        y = y1 + 72
        for row in rows:
            x = x1 + 25
            for item in row:
                draw_box(draw, (x, y, x + 120, y + 68), item, color, fnt=FONT_18)
                x += 135
            y += 100
    arrow(draw, (475, 270), (555, 270))
    arrow(draw, (955, 270), (1035, 270))
    arrow(draw, (690, 690), (880, 690))
    arrow(draw, (775, 555), (775, 410))
    draw.text((660, 450), "Airflow điều phối Spark, training và reload model", font=FONT_20, fill=COLORS["line"])
    return save(image, "11_deployment.png")


def draw_ui_wireframe() -> Path:
    image, draw = new_canvas("Thiết kế các giao diện chính")
    screens = [
        ((70, 135, 500, 820), "Login", ["Hospital Readmission MLOps", "username", "password", "Doctor / Admin buttons"]),
        ((585, 135, 1015, 820), "Admin Console", ["Dashboard stats", "User accounts", "Disable / enable user", "MLOps tools"]),
        ((1100, 135, 1530, 820), "Patient & Prediction", ["Patient directory", "Clinical form", "Prediction result", "High-risk list"]),
    ]
    for box, title, rows in screens:
        draw.rounded_rectangle(box, radius=22, fill="#f7f7f7", outline=COLORS["dark_blue"], width=4)
        draw.rectangle((box[0], box[1], box[2], box[1] + 58), fill=COLORS["dark_blue"])
        draw.text((box[0] + 20, box[1] + 16), title, font=FONT_22, fill="#ffffff")
        y = box[1] + 95
        for idx, row in enumerate(rows):
            color = [COLORS["blue"], COLORS["green"], COLORS["orange"], COLORS["purple"]][idx % 4]
            draw_box(draw, (box[0] + 40, y, box[2] - 40, y + 75), row, color, fnt=FONT_18)
            y += 105
        draw.rounded_rectangle((box[0] + 40, box[3] - 115, box[2] - 40, box[3] - 55), radius=12, fill="#eeeeee", outline="#808080")
        draw.text((box[0] + 65, box[3] - 98), "responsive React + Vite + lucide icons", font=FONT_18, fill=COLORS["line"])
    return save(image, "12_ui_wireframe.png")


def draw_monitoring() -> Path:
    image, draw = new_canvas("Sơ đồ quan sát và cảnh báo vận hành")
    boxes = [
        ((90, 200, 390, 285), "FastAPI /metrics", COLORS["green"]),
        ((520, 200, 820, 285), "Prometheus scrape\n5 giây/lần", COLORS["orange"]),
        ((950, 200, 1250, 285), "Grafana dashboard", COLORS["purple"]),
        ((950, 500, 1250, 585), "Alert rules", COLORS["red"]),
        ((90, 500, 390, 585), "Metrics\nrequests, latency, errors,\nrisk_level, model_reload", COLORS["blue"]),
    ]
    for box, label, color in boxes:
        draw_box(draw, box, label, color, fnt=FONT_18)
    arrow(draw, (390, 242), (520, 242))
    arrow(draw, (820, 242), (950, 242))
    arrow(draw, (670, 285), (1075, 500), fill=COLORS["red"])
    arrow(draw, (240, 500), (240, 285), fill=COLORS["blue"])
    alerts = [
        "FastAPIDown: up == 0 trong 1 phút",
        "PredictionErrorsDetected: tăng lỗi trong 5 phút",
        "HighPredictionLatency: p95 > 2 giây",
        "ModelReloadFailures: reload model thất bại",
    ]
    y = 690
    for alert_text in alerts:
        draw.text((360, y), alert_text, font=FONT_20, fill=COLORS["line"])
        y += 42
    return save(image, "13_monitoring.png")


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)
    for rel_id, rel in list(doc.part.rels.items()):
        if "image" in rel.reltype:
            del doc.part.rels[rel_id]


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color.replace("#", ""))
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        shade_cell(table.rows[0].cells[idx], "D9D9D9")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph("")


def style_paragraph(paragraph, style_name: str | None = None, align=None) -> None:
    if style_name:
        try:
            paragraph.style = style_name
        except KeyError:
            pass
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)


def add_p(doc: Document, text: str, style_name: str = "Nội dung văn bản", align=None):
    p = doc.add_paragraph(text)
    style_paragraph(p, style_name, align)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_caption(doc: Document, text: str, kind: str = "figure") -> None:
    style = "FigureCaption" if kind == "figure" else "TableCaption"
    p = doc.add_paragraph(text)
    style_paragraph(p, style, WD_ALIGN_PARAGRAPH.CENTER)


def add_diagram(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.4))
    add_caption(doc, caption, "figure")


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        try:
            p = doc.add_paragraph(item, style="List Bullet")
        except KeyError:
            p = doc.add_paragraph(f"- {item}")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_center_line(
    doc: Document,
    text: str = "",
    size: int = 14,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)


def add_cover(doc: Document, logo_path: Path) -> None:
    add_center_line(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP TP HỒ CHÍ MINH", 14, True)
    add_center_line(doc, "KHOA CÔNG NGHỆ THÔNG TIN", 14, True)
    add_center_line(doc)
    add_center_line(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo_path), width=Inches(4.16))

    add_center_line(doc)
    add_center_line(doc, "ĐỒ ÁN CUỐI KÌ", 16, True)
    add_center_line(doc, "CÔNG NGHỆ MỚI TRONG PHÁT TRIỂN ỨNG DỤNG", 16, True)
    add_center_line(doc)
    add_center_line(doc, "HỆ THỐNG MLOPS DỰ ĐOÁN", 16, True)
    add_center_line(doc, "TÁI NHẬP VIỆN BỆNH NHÂN", 16, True)

    for _ in range(6):
        add_center_line(doc)

    add_center_line(
        doc,
        "Người thực hiện:   TRẦN THÁI HÀ – 22636801",
        14,
        False,
        WD_ALIGN_PARAGRAPH.RIGHT,
    )
    add_center_line(doc, "Lớp       :    DHKHDL18A", 14, False, WD_ALIGN_PARAGRAPH.RIGHT)
    add_center_line(doc, "Khoá     :    18", 14, False, WD_ALIGN_PARAGRAPH.RIGHT)
    add_center_line(
        doc,
        "Người hướng dẫn: TS. BÙI THANH HÙNG",
        14,
        False,
        WD_ALIGN_PARAGRAPH.RIGHT,
    )

    for _ in range(3):
        add_center_line(doc)

    add_center_line(doc, "THÀNH PHỐ HỒ CHÍ MINH, NĂM 2026", 14, True)
    doc.add_page_break()


def metrics_rows() -> list[list[str]]:
    data = json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    rows = []
    for name, value in data["metrics"].items():
        rows.append([name, f"{float(value):.4f}", "Chỉ số đánh giá mô hình champion"])
    return rows


def candidate_rows() -> list[list[str]]:
    data = json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    rows = []
    for candidate in data["all_candidates"]:
        m = candidate["metrics"]
        rows.append([
            candidate["candidate_name"],
            candidate["model_type"],
            f"{m['accuracy']:.4f}",
            f"{m['precision']:.4f}",
            f"{m['recall']:.4f}",
            f"{m['f1_score']:.4f}",
            f"{m['roc_auc']:.4f}",
        ])
    return rows


def build_doc() -> None:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    for old_diagram in DIAGRAM_DIR.glob("*.png"):
        old_diagram.unlink()

    logo_path = extract_logo()

    diagrams = [
        draw_function_diagram(),
        draw_use_case(),
        prepare_source_diagram("activity", "03_activity.png"),
        prepare_source_diagram("sequence", "04_sequence.png"),
        draw_class_diagram(),
        draw_erd(),
        draw_dfd(),
        draw_ui_wireframe(),
        draw_algorithm(),
        draw_retraining(),
        draw_mlops_diagram(),
        draw_deployment(),
        draw_monitoring(),
    ]

    doc = Document(str(TEMPLATE))
    clear_document(doc)
    add_cover(doc, logo_path)

    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_p(doc, "Trước tiên, em xin gửi lời cảm ơn chân thành đến TS. Bùi Thanh Hùng, người đã định hướng và hỗ trợ em trong quá trình thực hiện đồ án cuối kì môn Công nghệ mới trong phát triển ứng dụng.")
    add_p(doc, "Em xin cảm ơn quý thầy cô Khoa Công nghệ Thông tin, Trường Đại học Công nghiệp Thành phố Hồ Chí Minh, đã trang bị các kiến thức nền tảng về phát triển phần mềm, dữ liệu, trí tuệ nhân tạo và triển khai hệ thống.")
    add_p(doc, "Trong quá trình thực hiện, em đã cố gắng xây dựng một hệ thống hoàn chỉnh theo hướng MLOps. Tuy nhiên, đồ án khó tránh khỏi thiếu sót, em rất mong nhận được góp ý của quý thầy cô để sản phẩm được hoàn thiện hơn.")
    add_p(doc, "Em xin chân thành cảm ơn!")
    add_p(doc, "TP. Hồ Chí Minh, ngày   tháng   năm", "Normal", WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "Tác giả", "Normal", WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "Trần Thái Hà", "Normal", WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()

    add_p(doc, "ĐỒ ÁN ĐƯỢC HOÀN THÀNH", "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "TẠI TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP TP HỒ CHÍ MINH", "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Tôi xin cam đoan đây là sản phẩm đồ án của riêng tôi và được sự hướng dẫn của TS. Bùi Thanh Hùng. Các nội dung nghiên cứu, kết quả trong đề tài là trung thực và được xây dựng dựa trên mã nguồn, dữ liệu và cấu hình hiện có trong project Hospital_Readmission_Prediction.")
    add_p(doc, "Những tài liệu, thư viện và công nghệ tham khảo trong đồ án đều được nêu ở phần tài liệu tham khảo. Nếu phát hiện có bất kỳ sai sót nào, tôi xin chịu trách nhiệm về nội dung báo cáo.")
    add_p(doc, "TP. Hồ Chí Minh, ngày   tháng   năm", "Normal", WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "Tác giả", "Normal", WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "Trần Thái Hà", "Normal", WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()

    add_heading(doc, "PHẦN ĐÁNH GIÁ CỦA GIẢNG VIÊN", 1)
    for _ in range(7):
        add_p(doc, "_" * 115, "Normal")
    add_p(doc, "Tp. Hồ Chí Minh, ngày     tháng   năm", "Normal", WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "(kí và ghi họ tên)", "Normal", WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()

    add_heading(doc, "TÓM TẮT", 1)
    add_p(doc, "Tái nhập viện trong thời gian ngắn là một vấn đề quan trọng trong quản trị bệnh viện vì làm tăng chi phí điều trị, gây áp lực lên nhân lực y tế và phản ánh nguy cơ biến chứng sau xuất viện. Đồ án này xây dựng hệ thống MLOps dự đoán khả năng bệnh nhân tiểu đường tái nhập viện trong vòng dưới 30 ngày dựa trên dữ liệu lâm sàng, thông tin nhập viện, chẩn đoán và thuốc điều trị.")
    add_p(doc, "Về dữ liệu, project sử dụng tệp diabetic_data.csv gồm 101.766 lượt khám với 50 cột. Nhãn nhị phân được tạo từ cột readmitted, trong đó giá trị <30 được xem là ca tái nhập viện sớm. Pipeline xử lý dữ liệu loại bỏ một số cột thiếu nhiều hoặc ít hữu ích, chuẩn hóa giá trị thiếu, mã hóa biến phân loại và huấn luyện các mô hình ứng viên bằng scikit-learn.")
    add_p(doc, "Về hệ thống, đồ án triển khai đầy đủ các thành phần MLOps: Spark batch tạo dữ liệu gold và offline feature, Kafka/Spark Streaming là luồng ingestion tùy chọn, Airflow điều phối ingestion/ETL/training/DB-triggered retraining, MLflow và MinIO quản lý thí nghiệm cùng artifact mô hình, FastAPI phục vụ suy luận, PostgreSQL và Redis lưu log, React cung cấp giao diện nghiệp vụ, Prometheus và Grafana giám sát vận hành.")
    add_p(doc, "Kết quả hiện tại cho thấy Random Forest là mô hình champion, đạt ROC-AUC khoảng 0,648 trên tập kiểm thử. Dù chất lượng mô hình còn có thể cải thiện, hệ thống đã thể hiện được quy trình đầu cuối từ dữ liệu, huấn luyện, đăng ký mô hình, phục vụ dự đoán đến giám sát và tái huấn luyện.")
    doc.add_page_break()

    add_heading(doc, "MỤC LỤC", 1)
    toc = [
        "LỜI CẢM ƠN",
        "PHẦN ĐÁNH GIÁ CỦA GIẢNG VIÊN",
        "TÓM TẮT",
        "DANH MỤC KÍ HIỆU VÀ CHỮ VIẾT TẮT",
        "DANH MỤC CÁC HÌNH VẼ",
        "DANH MỤC CÁC BẢNG",
        "CHƯƠNG 1 – PHÂN TÍCH, THIẾT KẾ",
        "1.1. Mô tả bài toán",
        "1.2. Sơ đồ chức năng tổng quát",
        "1.3. Biểu đồ trường hợp sử dụng",
        "1.4. Biểu đồ hoạt động",
        "1.5. Biểu đồ trình tự",
        "1.6. Biểu đồ dữ liệu và lớp",
        "1.7. Biểu đồ luồng dữ liệu",
        "1.8. Thiết kế giao diện",
        "1.9. Thiết kế giải thuật",
        "1.10. Thiết kế kiểm thử",
        "CHƯƠNG 2 – HIỆN THỰC",
        "2.1. Công nghệ sử dụng",
        "2.2. Các chức năng đã hiện thực",
        "2.3. Kết quả đánh giá mô hình",
        "CHƯƠNG 3 – KẾT LUẬN",
        "TÀI LIỆU THAM KHẢO",
        "PHỤ LỤC",
    ]
    for item in toc:
        add_p(doc, item, "Normal")
    doc.add_page_break()

    add_heading(doc, "DANH MỤC KÍ HIỆU VÀ CHỮ VIẾT TẮT", 1)
    abbreviations = [
        ("AI", "Artificial Intelligence – Trí tuệ nhân tạo"),
        ("API", "Application Programming Interface – Giao diện lập trình ứng dụng"),
        ("DAG", "Directed Acyclic Graph – Đồ thị có hướng không chu trình"),
        ("DVC", "Data Version Control – Quản lý phiên bản dữ liệu"),
        ("ETL", "Extract, Transform, Load – Trích xuất, biến đổi và nạp dữ liệu"),
        ("JWT", "JSON Web Token – Chuẩn token xác thực"),
        ("MLOps", "Machine Learning Operations – Vận hành vòng đời mô hình học máy"),
        ("MLflow", "Nền tảng theo dõi thí nghiệm và quản lý model registry"),
        ("ROC-AUC", "Diện tích dưới đường cong ROC"),
        ("UI", "User Interface – Giao diện người dùng"),
    ]
    for key, value in abbreviations:
        add_p(doc, f"{key}\t\t{value}", "Normal")
    doc.add_page_break()

    figure_titles = [
        "Hình 1.1: Sơ đồ chức năng tổng quát của hệ thống",
        "Hình 1.2: Biểu đồ trường hợp sử dụng",
        "Hình 1.3: Biểu đồ hoạt động dự đoán và phản hồi",
        "Hình 1.4: Biểu đồ trình tự xử lý dự đoán",
        "Hình 1.5: Biểu đồ lớp của hệ thống",
        "Hình 1.6: Biểu đồ quan hệ dữ liệu PostgreSQL",
        "Hình 1.7: Biểu đồ luồng dữ liệu batch, streaming tùy chọn và serving",
        "Hình 1.8: Thiết kế các giao diện chính",
        "Hình 1.9: Quy trình tiền xử lý, huấn luyện và lựa chọn mô hình",
        "Hình 1.10: Luồng tái huấn luyện tự động",
        "Hình 1.11: Sơ đồ MLOps tổng thể",
        "Hình 2.1: Sơ đồ triển khai Docker Compose",
        "Hình 2.2: Sơ đồ quan sát và cảnh báo vận hành",
    ]
    add_heading(doc, "DANH MỤC CÁC HÌNH VẼ", 1)
    for title in figure_titles:
        add_p(doc, title, "Normal")
    doc.add_page_break()

    table_titles = [
        "Bảng 1.1: Các trường hợp sử dụng chính của hệ thống",
        "Bảng 1.2: Các bảng dữ liệu chính",
        "Bảng 1.3: Các màn hình chính của giao diện",
        "Bảng 1.4: Các loại kiểm thử",
        "Bảng 2.1: Công nghệ sử dụng trong hệ thống",
        "Bảng 2.2: So sánh mô hình ứng viên",
        "Bảng 2.3: Kết quả đánh giá mô hình champion",
        "Bảng P.1: Các địa chỉ truy cập dịch vụ",
        "Bảng P.2: Các endpoint API chính",
    ]
    add_heading(doc, "DANH MỤC CÁC BẢNG", 1)
    for title in table_titles:
        add_p(doc, title, "Normal")
    doc.add_page_break()

    add_heading(doc, "CHƯƠNG 1 – PHÂN TÍCH, THIẾT KẾ", 1)
    add_heading(doc, "1.1. Mô tả bài toán", 2)
    add_p(doc, "Bài toán của đồ án là dự đoán nguy cơ tái nhập viện sớm của bệnh nhân tiểu đường sau một lượt điều trị nội trú. Trong dữ liệu, nhãn readmitted có ba giá trị: NO, >30 và <30. Đồ án chuyển bài toán thành phân loại nhị phân, trong đó readmitted = <30 được gán nhãn 1, còn NO hoặc >30 được gán nhãn 0.")
    add_p(doc, "Đầu vào của hệ thống là hồ sơ bệnh nhân gồm thông tin nhân khẩu học cơ bản, thông tin lần nhập viện, số ngày nằm viện, số xét nghiệm, số thủ thuật, số thuốc, lịch sử khám ngoại trú/cấp cứu/nội trú, nhóm chẩn đoán và trạng thái sử dụng thuốc điều trị tiểu đường. Đầu ra là xác suất tái nhập viện, nhãn dự đoán và mức rủi ro low, medium hoặc high.")
    add_p(doc, "Mục tiêu của đồ án không chỉ là huấn luyện mô hình, mà còn xây dựng một hệ thống vận hành mô hình theo chuẩn MLOps: dữ liệu được phiên bản hóa, pipeline huấn luyện có thể chạy lại, model được đăng ký trên MLflow, API phục vụ dự đoán có cơ chế reload model, log dự đoán được lưu lại để phân tích và kích hoạt tái huấn luyện.")

    add_heading(doc, "1.2. Sơ đồ chức năng tổng quát", 2)
    add_p(doc, "Sơ đồ chức năng tổng quát phân rã hệ thống dự đoán tái nhập viện bệnh nhân thành sáu nhóm chức năng chính. Cách phân chia này tách rõ phần nghiệp vụ dành cho bác sĩ, phần quản trị dành cho admin và phần vận hành MLOps phía sau hệ thống.")
    add_p(doc, "Nhóm quản lý người dùng và xác thực chịu trách nhiệm đăng nhập, cấp JWT token, phân quyền Doctor/Admin và cho phép người dùng xem thông tin tài khoản hiện tại. Đây là lớp kiểm soát truy cập chung trước khi người dùng thao tác với bệnh nhân, dự đoán hoặc các chức năng quản trị.")
    add_p(doc, "Nhóm quản lý bệnh nhân bao phủ các thao tác nghiệp vụ cốt lõi gồm thêm hồ sơ, xem danh sách, xem chi tiết, cập nhật, xóa hồ sơ bệnh nhân và xem danh sách bệnh nhân có nguy cơ cao. Dữ liệu bệnh nhân là đầu vào chính cho nhóm dự đoán tái nhập viện, nơi hệ thống nhận dữ liệu lâm sàng, tiền xử lý dữ liệu, tải model Production, trả kết quả dự đoán và ghi lại prediction log.")
    add_p(doc, "Nhóm quản trị MLOps cho phép admin theo dõi trạng thái pipeline, kích hoạt Airflow DAG, chạy ETL, huấn luyện model, retraining model và reload model Production. Nhóm quản lý dữ liệu và huấn luyện mô hình mô tả luồng xử lý offline: đọc dữ liệu nguồn, chạy Spark Batch ETL, tạo dữ liệu Gold, tạo Offline Feature, train và đánh giá model, ghi experiment vào MLflow, lưu artifact vào MinIO và promote model tốt nhất lên Production.")
    add_p(doc, "Nhóm giám sát hệ thống tập trung vào vận hành sau triển khai. FastAPI xuất metrics, Prometheus thu thập metrics, Grafana hiển thị dashboard và hệ thống có thể cảnh báo khi request, latency, error hoặc phân bố risk level xuất hiện bất thường. Nhờ đó, hệ thống không chỉ phục vụ dự đoán mà còn có khả năng theo dõi chất lượng vận hành và hỗ trợ tái huấn luyện khi dữ liệu mới phát sinh.")
    add_diagram(doc, diagrams[0], figure_titles[0])

    add_heading(doc, "1.3. Biểu đồ trường hợp sử dụng", 2)
    add_p(doc, "Hệ thống có hai tác nhân chính. Bác sĩ sử dụng hệ thống để quản lý bệnh nhân, nhập dữ liệu lâm sàng và xem kết quả dự đoán. Admin có thêm quyền xem toàn bộ dữ liệu, quản lý tài khoản người dùng, điều khiển pipeline MLOps, truy cập lịch sử dự đoán và reload model production.")
    add_diagram(doc, diagrams[1], figure_titles[1])
    add_caption(doc, "Bảng 1.1: Các trường hợp sử dụng chính của hệ thống", "table")
    add_table(doc, ["Tác nhân", "Trường hợp sử dụng", "Mô tả"], [
        ["Bác sĩ", "Đăng nhập", "Xác thực bằng username/password, nhận JWT để gọi API"],
        ["Bác sĩ", "Quản lý bệnh nhân", "Tạo, xem, cập nhật, xóa hồ sơ bệnh nhân thuộc phạm vi phụ trách"],
        ["Bác sĩ", "Dự đoán bệnh nhân", "Gửi đặc trưng lâm sàng tới API và nhận xác suất tái nhập viện"],
        ["Bác sĩ", "Xem high-risk", "Ưu tiên theo dõi bệnh nhân có xác suất tái nhập viện cao"],
        ["Admin", "Theo dõi dashboard", "Xem thống kê tổng quan, lịch sử dự đoán và thông tin model"],
        ["Admin", "Quản lý người dùng", "Tạo tài khoản, cập nhật họ tên/vai trò, reset mật khẩu và bật/tắt trạng thái hoạt động"],
        ["Admin", "Điều khiển pipeline", "Kích hoạt DAG ingestion, ETL, training và retraining"],
        ["Admin", "Reload model", "Yêu cầu FastAPI tải lại model Production từ MLflow Registry"],
    ])

    add_heading(doc, "1.4. Biểu đồ hoạt động", 2)
    add_p(doc, "Luồng hoạt động chính bắt đầu khi bác sĩ đăng nhập, tạo hoặc chọn hồ sơ bệnh nhân, sau đó gửi yêu cầu dự đoán. API chuẩn hóa dữ liệu đầu vào, gọi mô hình đã đăng ký trong MLflow, phân mức rủi ro và ghi lại kết quả. Khi hệ thống tích lũy đủ dữ liệu hoặc nhãn mới, Airflow có thể kích hoạt pipeline tái huấn luyện.")
    add_diagram(doc, diagrams[2], figure_titles[2])

    add_heading(doc, "1.5. Biểu đồ trình tự", 2)
    add_p(doc, "Biểu đồ trình tự mô tả tương tác giữa giao diện React, FastAPI, ModelLoader, MLflow Registry, Redis, PostgreSQL và Prometheus khi xử lý một dự đoán. Thiết kế này giúp API vừa trả kết quả nhanh cho người dùng, vừa lưu được log phục vụ dashboard, audit và tái huấn luyện.")
    add_diagram(doc, diagrams[3], figure_titles[3])

    add_heading(doc, "1.6. Biểu đồ dữ liệu và lớp", 2)
    add_p(doc, "Tầng dữ liệu quan hệ được hiện thực bằng PostgreSQL. Ở mức thiết kế lớp, các thành phần User, Patient, PredictionLog, RetrainingState, RetrainingRun và các service nghiệp vụ thể hiện trách nhiệm chính của hệ thống. Ở mức dữ liệu, các bảng cốt lõi gồm users, patients, prediction_logs, retraining_state và retraining_runs.")
    add_diagram(doc, diagrams[4], figure_titles[4])
    add_diagram(doc, diagrams[5], figure_titles[5])
    add_caption(doc, "Bảng 1.2: Các bảng dữ liệu chính", "table")
    add_table(doc, ["Bảng", "Vai trò trong hệ thống"], [
        ["users", "Lưu tài khoản, mật khẩu đã băm, họ tên, vai trò doctor/admin và trạng thái is_active"],
        ["patients", "Lưu hồ sơ bệnh nhân, đặc trưng lâm sàng và nhãn actual_readmitted nếu có"],
        ["prediction_logs", "Lưu request_json, prediction, probability, risk_level, model_name và thời điểm dự đoán"],
        ["retraining_state", "Ghi nhận số lượng prediction/patient đã dùng cho lần huấn luyện gần nhất"],
        ["retraining_runs", "Lưu lịch sử các lần tái huấn luyện, trigger_type, số bản ghi mới và trạng thái"],
    ])

    add_heading(doc, "1.7. Biểu đồ luồng dữ liệu", 2)
    add_p(doc, "Luồng dữ liệu chính của project là batch/offline. Dữ liệu CSV gốc được xử lý bằng Spark batch thành tầng gold và offline feature parquet, sau đó được dùng cho training và đăng ký model. Luồng Kafka/Spark Streaming vẫn được giữ dưới dạng tùy chọn để minh họa ingestion gần thời gian thực và ghi bronze parquet, nhưng không còn là trigger tái huấn luyện tự động.")
    add_diagram(doc, diagrams[6], figure_titles[6])

    add_heading(doc, "1.8. Thiết kế giao diện", 2)
    add_p(doc, "Giao diện được xây dựng bằng React và Vite. Các màn hình chính được thiết kế cho hai nhóm người dùng: bác sĩ thao tác với bệnh nhân và admin vận hành hệ thống MLOps. Admin có thêm màn hình Users để tạo tài khoản, cập nhật vai trò, reset mật khẩu và vô hiệu hóa/kích hoạt lại tài khoản. Giao diện sử dụng các khối thông tin ngắn, bảng danh sách và form lâm sàng để giảm thao tác lặp lại.")
    add_caption(doc, "Bảng 1.3: Các màn hình chính của giao diện", "table")
    add_table(doc, ["Màn hình", "Chức năng chính"], [
        ["Login", "Đăng nhập với tài khoản doctor01/doctor123 hoặc admin01/admin123"],
        ["Dashboard", "Hiển thị thống kê nghiệp vụ, thông tin model và Grafana dashboard MLOps"],
        ["Users", "Quản lý tài khoản doctor/admin, reset mật khẩu và bật/tắt trạng thái hoạt động"],
        ["Doctor Overview", "Tổng quan bệnh nhân của bác sĩ và danh sách ưu tiên"],
        ["Patients", "Quản lý danh bạ bệnh nhân, xem chi tiết, dự đoán và lịch sử từng bệnh nhân"],
        ["Create Patient", "Nhập hồ sơ bệnh nhân mới và chạy dự đoán ngay sau khi tạo"],
        ["High Risk Patients", "Danh sách bệnh nhân có nguy cơ tái nhập viện cao"],
        ["Pipelines", "Xem trạng thái DAG và kích hoạt pipeline Airflow"],
        ["Observability", "Truy cập Prometheus, MLflow và Airflow trong giao diện quản trị"],
    ])
    add_diagram(doc, diagrams[7], figure_titles[7])

    add_heading(doc, "1.9. Thiết kế giải thuật", 2)
    add_p(doc, "Pipeline huấn luyện đọc dữ liệu CSV, thay ký tự ? bằng missing value, tạo nhãn readmitted_binary, loại bỏ các cột encounter_id, patient_nbr, weight, payer_code, medical_specialty và readmitted. Sau đó dữ liệu được chia train/test theo tỷ lệ 80/20 có stratify để giữ phân bố nhãn.")
    add_p(doc, "Mô hình được xây dựng bằng sklearn Pipeline. Các cột số đi qua SimpleImputer(strategy='median') và StandardScaler; các cột phân loại đi qua SimpleImputer(fill_value='Unknown') và OneHotEncoder(handle_unknown='ignore'). Hai ứng viên trong cấu hình hiện tại là Random Forest và Logistic Regression, champion được chọn theo recall.")
    add_diagram(doc, diagrams[8], figure_titles[8])
    add_p(doc, "Cơ chế tái huấn luyện được thiết kế theo hướng database-triggered. Khi đủ bệnh nhân hoặc prediction log đã có nhãn actual_readmitted trong PostgreSQL, Airflow chạy prepare_training_from_db để tạo tập db_training_data.csv, huấn luyện lại hai mô hình ứng viên, chọn candidate tốt nhất theo recall, đăng ký phiên bản mới với nhãn v1, v2... và gọi FastAPI reload model.")
    add_diagram(doc, diagrams[9], figure_titles[9])
    add_p(doc, "Sơ đồ MLOps tổng thể mô tả vòng đời khép kín của mô hình: dữ liệu được thu thập và xử lý, thí nghiệm được ghi nhận trong MLflow, model tốt nhất được đưa vào Registry, FastAPI phục vụ suy luận, kết quả vận hành được giám sát và dữ liệu mới quay lại pipeline tái huấn luyện.")
    add_diagram(doc, diagrams[10], figure_titles[10])

    add_heading(doc, "1.10. Thiết kế kiểm thử", 2)
    add_p(doc, "Kiểm thử hệ thống cần bao phủ cả logic học máy, API nghiệp vụ và vận hành MLOps. Vì hệ thống có nhiều service, kiểm thử tích hợp đặc biệt quan trọng để bảo đảm dữ liệu đi đúng từ giao diện, API, cơ sở dữ liệu đến dashboard giám sát.")
    add_caption(doc, "Bảng 1.4: Các loại kiểm thử", "table")
    add_table(doc, ["Loại kiểm thử", "Phạm vi", "Ví dụ"], [
        ["Unit test", "Hàm tiền xử lý, phân mức rủi ro, normalize patient data", "Xác suất 0,75 phải trả về risk_level = high"],
        ["API test", "Endpoint FastAPI và phân quyền JWT", "Chưa đăng nhập gọi /prediction-logs phải bị từ chối"],
        ["Model test", "Chất lượng mô hình và tính tái lập", "reports/metrics.json có đủ accuracy, precision, recall, f1_score, roc_auc"],
        ["Integration test", "React - FastAPI - MLflow - PostgreSQL", "Tạo bệnh nhân, dự đoán, kiểm tra log xuất hiện trong prediction_logs"],
        ["Pipeline test", "Airflow DAG và Spark job", "Chạy etl_dag tạo data/gold và data/features thành công"],
        ["Monitoring test", "Prometheus/Grafana", "FastAPI expose đủ metric prediction_requests_total và prediction_errors_total"],
    ])

    add_heading(doc, "CHƯƠNG 2 – HIỆN THỰC", 1)
    add_heading(doc, "2.1. Công nghệ sử dụng", 2)
    add_p(doc, "Project được đóng gói bằng Docker Compose, gồm các service phục vụ dữ liệu, huấn luyện mô hình, API, giao diện và giám sát. Việc triển khai bằng container giúp tái lập môi trường và giảm sai khác giữa máy phát triển và môi trường chạy thử.")
    add_caption(doc, "Bảng 2.1: Công nghệ sử dụng trong hệ thống", "table")
    add_table(doc, ["Nhóm", "Công nghệ", "Vai trò"], [
        ["Frontend", "React 18, Vite, axios, lucide-react", "Giao diện bác sĩ/admin và gọi FastAPI"],
        ["Backend", "FastAPI, Pydantic, JWT, passlib", "API dự đoán, xác thực, phân quyền, quản lý bệnh nhân"],
        ["Dữ liệu", "PostgreSQL, Redis", "Lưu hồ sơ, prediction log, trạng thái retraining và cache log"],
        ["Streaming", "Kafka, Zookeeper, Spark Streaming", "Luồng tùy chọn để phát lại dữ liệu bệnh nhân và ghi bronze parquet"],
        ["Batch ETL", "Apache Spark", "Làm sạch dữ liệu, tạo gold và feature parquet"],
        ["ML/MLOps", "scikit-learn, MLflow, MinIO, DVC, Airflow", "Huấn luyện, theo dõi thí nghiệm, lưu artifact và điều phối pipeline"],
        ["Monitoring", "Prometheus, Grafana", "Thu thập metric, dashboard và alert rules"],
        ["Hạ tầng", "Docker Compose, Nginx", "Đóng gói service và reverse proxy cơ bản"],
    ])
    add_diagram(doc, diagrams[11], figure_titles[11])

    add_heading(doc, "2.2. Các chức năng đã hiện thực", 2)
    add_heading(doc, "2.2.1. Quản lý người dùng và phân quyền", 3)
    add_p(doc, "FastAPI khởi tạo hai tài khoản mặc định: doctor01/doctor123 và admin01/admin123. Người dùng đăng nhập qua /auth/login để nhận JWT. Các endpoint bệnh nhân yêu cầu vai trò doctor hoặc admin, trong khi endpoint pipeline, prediction log toàn hệ thống và quản trị user yêu cầu admin.")
    add_p(doc, "Admin có thể xem danh sách tài khoản, tạo user mới, cập nhật họ tên/vai trò, reset mật khẩu và thay đổi trạng thái is_active qua /users/{user_id}/status. Hệ thống không xóa cứng user; tài khoản bị vô hiệu hóa không thể đăng nhập và token cũ cũng bị từ chối khi gọi API.")
    add_heading(doc, "2.2.2. Quản lý bệnh nhân và dự đoán", 3)
    add_p(doc, "Người dùng có thể tạo, xem, cập nhật và xóa bệnh nhân. Endpoint /patients/{patient_id}/predict lấy dữ liệu bệnh nhân đã lưu, chuẩn hóa tên cột cho khớp pipeline huấn luyện, gọi model Production trên MLflow và lưu kết quả vào Redis cùng PostgreSQL.")
    add_heading(doc, "2.2.3. Pipeline ingestion, ETL và feature engineering", 3)
    add_p(doc, "ingestion_dag tạo Kafka topic patient-events và chạy producer phát lại dữ liệu CSV. etl_dag chạy Spark batch ETL để tạo data/gold/diabetic_gold.parquet, sau đó chạy feature_engineering.py để tạo data/features/offline/patient_features.parquet. Spark Streaming được đóng gói trong profile tùy chọn để ghi bronze parquet khi cần demo ingestion liên tục.")
    add_heading(doc, "2.2.4. Huấn luyện, đăng ký và reload model", 3)
    add_p(doc, "training.train huấn luyện các mô hình ứng viên, log params/metrics/model vào MLflow và lưu champion vào models/model.pkl cùng reports/metrics.json. training.register_model kiểm tra artifact, đăng ký champion của lần chạy hiện tại, gắn nhãn phiên bản v1, v2... và promote phiên bản mới lên Production.")
    add_heading(doc, "2.2.5. Giám sát và dashboard", 3)
    add_p(doc, "FastAPI expose /metrics theo chuẩn Prometheus, gồm tổng request dự đoán, latency, lỗi dự đoán, phân bố risk_level, histogram xác suất và số lần reload model. Grafana dashboard đọc Prometheus để hiển thị tình trạng API và chất lượng vận hành.")
    add_diagram(doc, diagrams[12], figure_titles[12])

    add_heading(doc, "2.3. Kết quả đánh giá mô hình", 2)
    add_p(doc, "Kết quả trong reports/metrics.json cho thấy candidate có recall cao nhất là champion hiện tại. Mô hình được chọn theo metric recall với mode max. Các chỉ số còn khiêm tốn, đặc biệt precision thấp, phản ánh độ khó của bài toán do nhãn <30 chỉ chiếm khoảng 11,2% dữ liệu.")
    add_caption(doc, "Bảng 2.2: So sánh mô hình ứng viên", "table")
    add_table(doc, ["Ứng viên", "Loại mô hình", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"], candidate_rows())
    add_caption(doc, "Bảng 2.3: Kết quả đánh giá mô hình champion", "table")
    add_table(doc, ["Chỉ số", "Giá trị", "Ý nghĩa"], metrics_rows())
    add_p(doc, "Kết quả ROC-AUC 0,6477 cho thấy mô hình đã học được tín hiệu phân biệt nhưng vẫn cần cải thiện trước khi dùng trong bối cảnh lâm sàng thật. Trong phiên bản hiện tại, trọng tâm của đồ án là hoàn thiện vòng đời MLOps đầu cuối; các hướng cải thiện mô hình được trình bày ở Chương 3.")

    add_heading(doc, "CHƯƠNG 3 – KẾT LUẬN", 1)
    add_heading(doc, "3.1. Kết quả đạt được", 2)
    add_p(doc, "Đồ án đã xây dựng được một hệ thống MLOps dự đoán tái nhập viện bệnh nhân tương đối đầy đủ, bao gồm dữ liệu, huấn luyện mô hình, phục vụ dự đoán, quản lý bệnh nhân, dashboard nghiệp vụ, pipeline tái huấn luyện và giám sát vận hành.")
    add_p(doc, "Về học máy, project đã hiện thực pipeline scikit-learn có tiền xử lý số/phân loại, thử nghiệm nhiều mô hình ứng viên, chọn champion theo ROC-AUC, log kết quả vào MLflow và lưu báo cáo metric. Về phần mềm, hệ thống có FastAPI, React, PostgreSQL, Redis, Kafka, Spark, Airflow, Prometheus và Grafana chạy trong Docker Compose.")
    add_heading(doc, "3.2. Hạn chế", 2)
    add_p(doc, "Mô hình hiện tại có ROC-AUC khoảng 0,648 và precision thấp, chưa đủ mạnh cho quyết định lâm sàng thực tế. Giao diện đã bao phủ các luồng chính nhưng vẫn chủ yếu phục vụ demo MLOps. Ngoài ra, hệ thống chưa có kiểm thử tự động đầy đủ cho toàn bộ API, DAG và frontend.")
    add_heading(doc, "3.3. Hướng phát triển trong tương lai", 2)
    add_p(doc, "Thứ nhất, cải thiện mô hình bằng cách thử thêm XGBoost/LightGBM, calibration xác suất, xử lý mất cân bằng lớp và tìm ngưỡng tối ưu theo mục tiêu recall hoặc F1. Thứ hai, bổ sung SHAP để giải thích yếu tố ảnh hưởng tới nguy cơ tái nhập viện. Thứ ba, hoàn thiện CI/CD, test tự động, kiểm soát schema dữ liệu và giám sát drift. Thứ tư, mở rộng triển khai sang Kubernetes khi cần scale nhiều bản sao service.")

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    references = [
        "[1] Apache Airflow Documentation, https://airflow.apache.org/docs/.",
        "[2] Apache Kafka Documentation, https://kafka.apache.org/documentation/.",
        "[3] Apache Spark Documentation, https://spark.apache.org/docs/latest/.",
        "[4] FastAPI Documentation, https://fastapi.tiangolo.com/.",
        "[5] Grafana Documentation, https://grafana.com/docs/.",
        "[6] MLflow Documentation, https://mlflow.org/docs/latest/.",
        "[7] Prometheus Documentation, https://prometheus.io/docs/.",
        "[8] scikit-learn Documentation, https://scikit-learn.org/stable/documentation.html.",
        "[9] DVC Documentation, https://dvc.org/doc.",
        "[10] UCI Machine Learning Repository, Diabetes 130-US hospitals for years 1999-2008 Data Set.",
    ]
    for ref in references:
        add_p(doc, ref)

    add_heading(doc, "PHỤ LỤC", 1)
    add_p(doc, "Phụ lục A. Hướng dẫn chạy hệ thống")
    add_p(doc, "Điều kiện cần: máy đã cài Docker và Docker Compose. Các bước cơ bản: sao chép .env.example thành .env nếu cần, bảo đảm data/diabetic_data.csv tồn tại, sau đó chạy docker compose up -d --build tại thư mục gốc project.")
    add_caption(doc, "Bảng P.1: Các địa chỉ truy cập dịch vụ", "table")
    add_table(doc, ["Thành phần", "Địa chỉ", "Tài khoản mặc định"], [
        ["Frontend", "http://localhost:3001", "doctor01/doctor123 hoặc admin01/admin123"],
        ["FastAPI", "http://localhost:8000", "JWT từ /auth/login"],
        ["FastAPI Docs", "http://localhost:8000/docs", "JWT"],
        ["MLflow", "http://localhost:5000/tools/mlflow/", "-"],
        ["Airflow", "http://localhost:8088/tools/airflow/", "admin/admin"],
        ["Kafka UI", "http://localhost:8080", "-"],
        ["MinIO Console", "http://localhost:9001", "minio/minio123"],
        ["Prometheus", "http://localhost:9090", "-"],
        ["Grafana", "http://localhost:3000", "admin/admin"],
    ])
    add_p(doc, "Phụ lục B. Cấu trúc thư mục chính")
    add_bullets(doc, [
        "airflow/: Dockerfile và các DAG ingestion, ETL, training, db-triggered retraining.",
        "spark/jobs/: Spark batch ETL, streaming ETL và feature engineering.",
        "training/: pipeline huấn luyện, chuẩn bị dữ liệu DB retraining và đăng ký model.",
        "inference/app/: FastAPI, auth, database, model loader, preprocessing và schema.",
        "frontend/src/: giao diện React cho dashboard, bệnh nhân, dự đoán, pipelines và observability.",
        "kafka/: producer, consumer và script tạo topic patient-events.",
        "monitoring/: cấu hình Prometheus, Grafana dashboard và alert rules.",
        "data/, models/, reports/: dữ liệu, model local và kết quả metric.",
    ])
    add_p(doc, "Phụ lục C. Các endpoint API chính")
    add_caption(doc, "Bảng P.2: Các endpoint API chính", "table")
    add_table(doc, ["Phương thức", "Đường dẫn", "Chức năng"], [
        ["POST", "/auth/login", "Đăng nhập và nhận access token"],
        ["GET", "/auth/me", "Lấy thông tin người dùng hiện tại"],
        ["GET/POST", "/users", "Admin xem danh sách hoặc tạo tài khoản doctor/admin"],
        ["PUT", "/users/{user_id}", "Admin cập nhật họ tên, vai trò hoặc reset mật khẩu"],
        ["PATCH", "/users/{user_id}/status", "Admin vô hiệu hóa hoặc kích hoạt lại tài khoản"],
        ["GET/POST", "/patients", "Danh sách hoặc tạo bệnh nhân"],
        ["GET/PUT/DELETE", "/patients/{patient_id}", "Xem, cập nhật hoặc xóa bệnh nhân"],
        ["POST", "/patients/{patient_id}/predict", "Dự đoán từ hồ sơ bệnh nhân đã lưu"],
        ["GET", "/patients/high-risk", "Danh sách bệnh nhân rủi ro cao"],
        ["POST", "/predict", "Dự đoán trực tiếp từ payload lâm sàng"],
        ["GET", "/prediction-logs", "Lịch sử dự đoán toàn hệ thống, dành cho admin"],
        ["GET", "/dashboard-stats", "Thống kê dashboard"],
        ["GET", "/mlops/pipelines", "Trạng thái các DAG Airflow"],
        ["POST", "/mlops/pipelines/{dag_id}/trigger", "Kích hoạt DAG được phép"],
        ["POST", "/reload-model", "Reload model Production từ MLflow"],
        ["GET", "/metrics", "Prometheus metrics"],
    ])

    doc.core_properties.title = "Hệ thống MLOps dự đoán tái nhập viện bệnh nhân"
    doc.core_properties.author = "Trần Thái Hà"
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build_doc()
